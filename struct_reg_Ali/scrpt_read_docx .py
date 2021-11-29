import xlsxwriter  # build excel,   sudo apt-get install python3-xlsxwriter
import docx

#Created by Emmanuel Carballo Vargas, and Marlon L. Coronado.


file_exc = xlsxwriter.Workbook('listed_info.xlsx')
sheet = file_exc.add_worksheet()

doc = docx.Document('preprocessed_doc.docx') #preprocessed doc in Google drive to delete any -section index-
tablesD = doc.tables

line_tex = []

for i in range(len(doc.paragraphs)): 

    line = doc.paragraphs[i].text
    line = line.split()

    if (len(line) == 3):

        line = [line[0], line[1]]
        line_tex.append(line)


i = 0
j = 0
k = 0


for table in tablesD:

    print('index = '+str(tablesD.index(table)))

    if tablesD.index(table)%2!=0: #Skip pair (un desired) tables, instead of delete them.

        if (len(table.rows) > 1): 

            k = k + 1

            sheet.write(i, 0, line_tex[k-1][0])
            sheet.write(i, 1, line_tex[k-1][1])

            noFirst = 0 #ommit the first row of every table...
            for row in table.rows:
            
                if noFirst!=0:

                    j = 0

                    for cell in row.cells:

                        #print(cell.text)
                        sheet.write(i, j+2, str(cell.text))

                        j = j + 1

                    i = i + 1

                noFirst = noFirst+1

file_exc.close()



