import xlsxwriter  # build excel,   sudo apt-get install python3-xlsxwriter
import docx

#Created by Emmanuel Carballo Vargas, and Marlon L. Coronado.


file_exc = xlsxwriter.Workbook('listed_info.xlsx')
sheet = file_exc.add_worksheet()


doc = docx.Document('prueba2.docx') 
tablesD = doc.tables


line_tex = [] #auxiliar lines vector to save info needed

for i in range(len(doc.paragraphs)): 

    line = doc.paragraphs[i].text #obtain the line text of the paragraph

    line = line.split() #split using 'space' as parameter

    #print(line)

    #conditions for the line to be saved. New: discard text -if this is not a 'NVMe H3' heading-
    if (len(line) == 2) and str(line[0])!='Table' and (doc.paragraphs[i].style.name=='NVMe H3'): 

        line = [line[0], line[1]]
        line_tex.append(line)




#using as indexing the table
i = 0  #index to run through rows of the exel...
j = 0  #index to run through the 2nd and further columns
k = 0  


for table in tablesD:

    #print('index = '+str(tablesD.index(table)))

    if i==0: #build 1st row of titles for the exel (can be omptimized)
        sheet.write(i, 0, 'Empty')
        sheet.write(i, 1, 'Description Header')
        sheet.write(i, 2, 'Base Bus')
        sheet.write(i, 3, 'Bits')
        sheet.write(i, 4, 'RW')
        sheet.write(i, 5, 'Description')
        sheet.write(i, 6, 'Default')

        i=1

        continue #continue with the other rows of the exel



    if tablesD.index(table)%2!=0: #Skip pair (un desired) tables, instead of delete them.

        if (len(table.rows) > 1): 

            k = k + 1

            sheet.write(i, 0, '-')
            sheet.write(i, 1, line_tex[k-1][0])
            sheet.write(i, 2, line_tex[k-1][1].replace('(','').replace(')', ''))


            noFirst = 0 #ommit the first row of every table...
            for row in table.rows:
            
                if noFirst!=0:

                    j = 0

                    for cell in row.cells:

                        #print(cell.text)
                        sheet.write(i, j+3, str(cell.text))

                        j = j + 1

                    i = i + 1

                noFirst = noFirst+1

file_exc.close()



